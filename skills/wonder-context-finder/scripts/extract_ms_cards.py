#!/usr/bin/env python3
"""
Extract MS Cards from Word documents.

Requirements:
    pip install python-docx

Usage:
    python extract_ms_cards.py
    python extract_ms_cards.py --docs-dir "C:/path/to/docs" --output-dir "output"
"""
import re
import json
import argparse
from pathlib import Path

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. Run: pip install python-docx")

DEFAULT_DOCS_DIR = "C:/Users/44105/.claude/skills/wonder-context-finder/assets/source-docs"
DEFAULT_OUTPUT_DIR = "C:/Users/44105/.claude/skills/wonder-context-finder/references/ms-cards"

def extract_text_from_docx(docx_path: str) -> str:
    """Extract all text from a Word document."""
    if not DOCX_AVAILABLE:
        return ""

    doc = Document(docx_path)
    full_text = []

    for para in doc.paragraphs:
        full_text.append(para.text)

    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                row_text.append(cell.text)
            full_text.append(" | ".join(row_text))

    return "\n".join(full_text)

def find_ms_cards(text: str) -> list:
    """Find MS Card references in text."""
    # Pattern to match MS Card IDs like MS05-16, MS06-01, MS08-xx, etc.
    pattern = r'MS\d{2}-\d{1,2}'
    matches = re.findall(pattern, text)
    return list(set(matches))

def extract_card_section(text: str, card_id: str) -> str:
    """Extract the section for a specific MS Card."""
    # Look for the card ID followed by content until the next card or end
    pattern = rf'({card_id}[:\s].*?)(?=MS\d{{2}}-\d{{1,2}}[:\s]|$)'
    match = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
    if match:
        return match.group(1)[:5000]  # Limit content size
    return ""

def create_card_index(cards: dict) -> str:
    """Create an index markdown file for all MS Cards."""
    index_content = """# MS Cards Index

This index lists all MS Cards extracted from the Master Data V2.0 Use Cases documents.

## Overview

MS Cards document business requirements in the format `MSxx-yy`:
- **MS05 series**: Item detail UI cards
- **MS06 series**: BOM and component operations
- **MS08 series**: Major operations
- **MS13 series**: Complex features
- **MS15 series**: Additional operations

## Cards by Series

"""

    # Group cards by series
    series_groups = {}
    for card_id in sorted(cards.keys()):
        series = card_id[:4]  # MS05, MS06, etc.
        if series not in series_groups:
            series_groups[series] = []
        series_groups[series].append(card_id)

    for series in sorted(series_groups.keys()):
        index_content += f"### {series} Series\n\n"
        for card_id in sorted(series_groups[series]):
            card_info = cards[card_id]
            title = card_info.get('title', 'Unknown')
            index_content += f"- [{card_id}](cards/{card_id}.md) - {title}\n"
        index_content += "\n"

    return index_content

def main():
    parser = argparse.ArgumentParser(description='Extract MS Cards from Word documents')
    parser.add_argument('--docs-dir', default=DEFAULT_DOCS_DIR, help='Directory containing docx files')
    parser.add_argument('--output-dir', default=DEFAULT_OUTPUT_DIR, help='Output directory for extracted cards')
    parser.add_argument('--list-only', action='store_true', help='Only list found card IDs')

    args = parser.parse_args()

    docs_path = Path(args.docs_dir)
    output_path = Path(args.output_dir)

    if not DOCX_AVAILABLE:
        print("Error: python-docx is required. Install with: pip install python-docx")
        return

    if not docs_path.exists():
        print(f"Error: Documents directory not found: {docs_path}")
        return

    # Find all docx files
    docx_files = list(docs_path.glob("*.docx"))
    if not docx_files:
        print(f"No .docx files found in {docs_path}")
        return

    print(f"Found {len(docx_files)} document(s)")

    all_cards = {}
    all_text = ""

    for docx_file in docx_files:
        print(f"Processing: {docx_file.name}")
        text = extract_text_from_docx(str(docx_file))
        all_text += text + "\n"

        cards = find_ms_cards(text)
        print(f"  Found {len(cards)} card references: {', '.join(sorted(cards)[:10])}...")

        for card_id in cards:
            if card_id not in all_cards:
                all_cards[card_id] = {
                    'id': card_id,
                    'title': f'{card_id} Card',
                    'source_file': docx_file.name,
                    'content': ''
                }

    print(f"\nTotal unique cards found: {len(all_cards)}")

    if args.list_only:
        for card_id in sorted(all_cards.keys()):
            print(f"  {card_id}")
        return

    # Create output directories
    output_path.mkdir(parents=True, exist_ok=True)
    (output_path / "cards").mkdir(exist_ok=True)

    # Extract content for each card
    for card_id in all_cards:
        content = extract_card_section(all_text, card_id)
        all_cards[card_id]['content'] = content

        # Create individual card file
        card_file = output_path / "cards" / f"{card_id}.md"
        card_content = f"""# {card_id}

**Source**: Master Data V2.0 Use Cases

## Content

{content if content else "Content extraction pending. Please refer to the original document."}

## Related Information

- Search Jira: `mcp__atlassian__search(query="{card_id}")`
- Search git: `git log --all --oneline --grep="{card_id}"`
"""
        card_file.write_text(card_content, encoding='utf-8')

    # Create index file
    index_content = create_card_index(all_cards)
    (output_path / "index.md").write_text(index_content, encoding='utf-8')

    print(f"\nExtracted {len(all_cards)} cards to {output_path}")
    print(f"Index file: {output_path / 'index.md'}")

if __name__ == "__main__":
    main()
